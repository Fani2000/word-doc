from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading("Testing the python-docx module.")

name = input("Enter you name: ")

p = document.add_paragraph(f"Here is the content of the template.....Name: ")
p.add_run(name).bold = True

document.save('test.docx')

